#!/usr/bin/env python3
"""Insert experiment data tables into the thesis document at appropriate positions."""
import csv, os, copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = "/home/ubuntu/pythonProject1/张延开毕业论文修改版.docx"
OUT = "/home/ubuntu/pythonProject1/张延开毕业论文修改版.docx"
RES = "RESULT/fine_experiments"

doc = Document(DOC)

# ── style helpers ────────────────────────────────────────────────────
def set_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def insert_para_after(doc, paragraph, text='', style=None):
    """Insert a paragraph *after* the given paragraph. Returns the new Paragraph."""
    next_elem = paragraph._element.getnext()
    if next_elem is not None and next_elem.tag == qn('w:p'):
        # Find the Paragraph wrapper and insert before it (= after our target)
        for p in doc.paragraphs:
            if p._element is next_elem:
                return p.insert_paragraph_before(text, style)
    # No next paragraph — append to end
    return doc.add_paragraph(text, style=style)


def add_table_and_caption(doc, headers, rows, insert_after_idx, caption_text, table_num):
    """Add table + caption after paragraph index, returns None."""
    full_caption = f"表 3.3-{table_num} {caption_text}"

    cap_p = doc.paragraphs[insert_after_idx]
    new_cap = insert_para_after(doc, cap_p, full_caption, doc.styles['Caption'])

    # Create table
    n_cols = len(headers)
    n_rows = len(rows) + 1
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for j, h in enumerate(headers):
        set_cell(tbl.rows[0].cells[j], h, bold=True, size=9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell(tbl.rows[i+1].cells[j], val, bold=False, size=8)

    # Insert table after caption
    new_cap._element.addnext(tbl._tbl)


# ══════════════════════════════════════════════════════════════════════
#  Read data from CSV files
# ══════════════════════════════════════════════════════════════════════

# 1. δ sensitivity (exp1)
deltas = []
with open(f"{RES}/exp1_delta_sensitivity.csv") as f:
    for row in csv.DictReader(f):
        deltas.append((row['delta'], row['rate']))

# 2. JPEG attack (exp4)
jpeg_data = []
with open(f"{RES}/exp4_attack_robustness.csv") as f:
    for row in csv.DictReader(f):
        if row['attack'] == 'jpeg':
            jpeg_data.append((row['param'], row['rate']))

# 3. Gaussian noise (exp4)
gauss_data = []
with open(f"{RES}/exp4_attack_robustness.csv") as f:
    for row in csv.DictReader(f):
        if row['attack'] == 'gaussian':
            gauss_data.append((row['param'], row['rate']))

# 4. Crop (exp4)
crop_data = []
with open(f"{RES}/exp4_attack_robustness.csv") as f:
    for row in csv.DictReader(f):
        if row['attack'] == 'crop':
            crop_data.append((row['param'], row['rate']))

# 5. ECC config (exp2)
ecc_data = []
with open(f"{RES}/exp2_ecc_length.csv") as f:
    for row in csv.DictReader(f):
        ecc_data.append((row['msg_len'], row['ecc_len'], row['total_bits'], row['rate']))

# 6. Method comparison (rebuttal/expB)
method_data = []
with open(f"{RES}/rebuttal/expB_method.csv") as f:
    for row in csv.DictReader(f):
        method_data.append((row['delta'], row['ref_rate'], row['inv_rate']))

# 7. PSNR/SSIM (rebuttal/expA)
psnr_data = []
with open(f"{RES}/rebuttal/expA_quality.csv") as f:
    for row in csv.DictReader(f):
        psnr_data.append((row['delta'], row['latent_psnr'], row['pixel_psnr'], row['ssim']))


# ══════════════════════════════════════════════════════════════════════
#  Insert tables (bottom-up to avoid index shifting)
# ══════════════════════════════════════════════════════════════════════
# Each call inserts 1 caption paragraph before the next paragraph.
# Processing in reverse (highest index first) means earlier insertions
# don't shift the targets of later insertions.

# ── 7. Method comparison → after Para 238 (3.3.5) ──
add_table_and_caption(doc,
    ['δ'] + [str(r[0]) for r in method_data],
    [['Reference (%)'] + [r[1] for r in method_data],
     ['DDIM Inversion (%)'] + [r[2] for r in method_data]],
    238,
    f"Extraction Method: Reference vs DDIM Inversion (6 images per δ)", 7)

# ── 6. PSNR/SSIM → after Para 228 (empty, before delta figures in 3.3.4) ──
add_table_and_caption(doc,
    ['δ'] + [str(r[0]) for r in psnr_data],
    [['Latent PSNR (dB)'] + [r[1] for r in psnr_data],
     ['Pixel PSNR (dB)'] + [r[2] for r in psnr_data],
     ['SSIM'] + [r[3] for r in psnr_data]],
    228,
    f"Image Quality Metrics vs Embedding Strength (6 images per δ)", 6)

# ── 5. ECC config → after Para 231 (in 3.3.4) ──
add_table_and_caption(doc,
    ['MSG_LEN', 'ECC_LEN', 'Total Bits', 'Success Rate (%)'],
    [[r[0], r[1], r[2], r[3]] for r in ecc_data],
    231,
    f"Reed-Solomon Code Configuration (δ=2.0, 12 images per config)", 5)

# ── 4. Random crop → after Para 221 (Normal with crop results) ──
add_table_and_caption(doc,
    ['Retention'] + [f"{float(r[0]):.0%}" for r in crop_data],
    [['Rate (%)'] + [r[1] for r in crop_data]],
    221,
    f"Random Crop Attack (δ=2.8, 60 images per retention)", 4)

# ── 3. Gaussian noise → after Para 216 (Normal with Gaussian results) ──
add_table_and_caption(doc,
    ['σ'] + [str(r[0]) for r in gauss_data],
    [['Rate (%)'] + [r[1] for r in gauss_data]],
    216,
    f"Gaussian Noise Attack (δ=2.8, 60 images per σ)", 3)

# ── 2. JPEG attack → after Para 209 (Normal with JPEG results) ──
add_table_and_caption(doc,
    ['Quality'] + [r[0] for r in jpeg_data],
    [['Rate (%)'] + [r[1] for r in jpeg_data]],
    209,
    f"JPEG Compression Attack (δ=2.8, 60 images per quality)", 2)

# ── 1. δ sensitivity → after Para 206 (3.2 秘密提取准确性评估) ──
add_table_and_caption(doc,
    ['δ'] + [r[0] for r in deltas],
    [['Success Rate (%)'] + [r[1] for r in deltas]],
    206,
    f"δ Sensitivity (Dense Sampling, 15 images per δ)", 1)


# ══════════════════════════════════════════════════════════════════════
#  Save
# ══════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
print("Inserted 7 tables total")
