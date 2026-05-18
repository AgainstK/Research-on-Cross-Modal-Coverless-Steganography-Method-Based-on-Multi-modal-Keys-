#!/usr/bin/env python3
"""Fill in the graduation thesis guidance record form."""
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/ubuntu/pythonProject1/RESULT/fine_experiments/1.指导记录表(附件3：暨南大学本科毕业论文（设计）指导记录表).docx"
OUT = "/home/ubuntu/pythonProject1/RESULT/fine_experiments/1.指导记录表(附件3：暨南大学本科毕业论文（设计）指导记录表)（已填写）.docx"

doc = Document(SRC)
tbl = doc.tables[0]

# ── Helper: set cell text while preserving formatting ──
def set_cell_text(cell, text):
    """Replace cell content with plain text."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    # Put text in first paragraph's first run
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

# ── Row 0: Student info ──
set_cell_text(tbl.rows[0].cells[1], '张延开')
set_cell_text(tbl.rows[0].cells[3], '2022101194')
set_cell_text(tbl.rows[0].cells[5], '网络空间安全')

# ── Row 1: Advisor info ──
set_cell_text(tbl.rows[1].cells[1], '李佩雅')
set_cell_text(tbl.rows[1].cells[3], '讲师')
set_cell_text(tbl.rows[1].cells[5], '暨南大学')

# ── Row 2: Title ──
set_cell_text(tbl.rows[2].cells[1], '基于多模态密钥的跨模态无载体隐写方法的研究')

# =====================================================================
#  Guidance records (each record spans 2 rows: time + content)
#  Template has 6 records (rows 3-14).
#  We need 11 records (every 2 weeks, Dec 1 → May 1)
#  So add 5 more records (10 rows) after the existing ones.
# =====================================================================

records = [
    {
        "date": "2025年12月1日",
        "content": (
            "本次指导确定了选题方向——基于扩散模型的无载体图像隐写方法研究。"
            "向学生介绍了该领域的研究背景和前沿动态，重点讲解了CRoSS等代表性工作的技术思路。"
            "要求学生广泛阅读相关文献，重点梳理扩散模型（DDPM、DDIM）的原理以及现有无载体隐写方法的优缺点，"
            "并在下次指导前提交文献调研报告和初步研究方案。"
        )
    },
    {
        "date": "2025年12月15日",
        "content": (
            "学生提交了文献调研报告，对扩散模型隐写的技术路线有了基本认识。"
            "讨论了开题报告的撰写框架，明确了研究目标和创新点：针对现有方法在密钥形式单一、"
            "跨模态信息隐藏不足的问题，拟采用多模态密钥（文本+深度图）融合策略，"
            "基于Stable Diffusion构建跨模态无载体隐写方案。"
            "对技术路线中的关键环节——确定性DDIM采样、潜空间信息嵌入、多条件控制——进行了详细论证，"
            "要求学生按照讨论结果完善开题报告。"
        )
    },
    {
        "date": "2026年1月5日",
        "content": (
            "审阅了开题报告初稿，整体框架合理，但存在以下问题：研究内容中的技术细节描述不够具体，"
            "实验方案中的评估指标和对比基准需进一步明确。建议学生在秘密信息嵌入策略部分补充潜空间扰动幅度的理论依据，"
            "并明确多模态密钥融合的具体实现方式（分阶段融合还是同步融合）。"
            "要求学生修改后提交正式版，准备参加开题答辩。"
        )
    },
    {
        "date": "2026年1月19日",
        "content": (
            "检查了实验环境搭建进度。学生在本地成功部署了Stable Diffusion v1.5和ControlNet，"
            "完成了预训练权重下载和Diffusers管道配置。讨论了技术方案的具体实现路径：考虑到实际算力限制，"
            "建议先以CPU模式进行小规模功能验证，再逐步迁移到GPU环境。"
            "明确了下一步的核心任务——实现CRoSS框架的基本隐写与提取流程，重点关注DDIM采样的可逆性保证。"
        )
    },
    {
        "date": "2026年2月2日",
        "content": (
            "学生实现了基于潜变量调制的隐写原型代码。在实现过程中遇到两个问题：一是秘密信息直接硬覆盖噪声维度导致生成图像质量退化，"
            "二是两阶段扩散时序控制失当导致公钥语义被覆盖。针对前者，建议采用小幅加性扰动策略替代硬覆盖；"
            "针对后者，建议重新分配两阶段步数比例（私钥阶段30%、公钥阶段70%）。"
            "修改后生成图像质量明显改善。要求学生继续完善提取模块，实现盲提取功能。"
        )
    },
    {
        "date": "2026年2月16日",
        "content": (
            "学生提交了初步实验结果。在δ敏感性实验中，发现δ≥2.7时提取成功率达100%，"
            "呈现明显的阈值效应。在鲁棒性实验中，JPEG压缩（质量因子≥75）和低强度高斯噪声下表现良好，"
            "但随机裁剪攻击下完全失效。建议补充Reed-Solomon纠错编码以增强鲁棒性，"
            "并对嵌入强度δ与图像质量、鲁棒性之间的权衡关系进行定量分析。要求设计ECC配置对比实验。"
        )
    },
    {
        "date": "2026年3月2日",
        "content": (
            "检查了ECC实验和LSB对比实验结果。RS编码在δ≥2.8时可将JPEG鲁棒性提升至质量因子40仍保持80%成功率。"
            "与LSB方法的对比表明本方法在抗JPEG压缩和高斯噪声方面具有显著优势。"
            "建议继续补充图像质量客观评估（PSNR/SSIM）实验，以定量说明不同δ取值对图像质量的影响，"
            "为δ选取提供理论依据。同时建议进行提取方法对比实验（参考比较法与DDIM反演法）。"
        )
    },
    {
        "date": "2026年3月16日",
        "content": (
            "学生提交了PSNR/SSIM质量评估和提取方法对比实验结果。结果表明δ=2.8时潜空间PSNR为8.75dB，"
            "像素域PSNR为13.03dB，SSIM为0.312。参考比较法在δ≥2.8时达到100%成功率，"
            "而DDIM反演法在所有δ下均为0%。建议在论文中重点阐述潜空间PSNR这一客观质量指标的意义，"
            "并说明扩散模型的混沌敏感性导致像素域PSNR/SSIM偏低但视觉质量良好的原因。"
            "要求学生开始撰写论文初稿。"
        )
    },
    {
        "date": "2026年4月1日",
        "content": (
            "审阅了论文初稿的前三章。整体结构和逻辑框架基本合理，但仍需改进以下几个方面："
            "一是摘要和引言中的研究动机表述需更加凝练，突出跨模态隐写的挑战和本文的创新点；"
            "二是第二章技术原理部分应对扩散模型和DDIM采样的数学推导进行适当展开；"
            "三是实验部分的图表规范需统一，表格应使用三线表格式，图中字体需清晰可读。"
            "要求学生在一周内完成修改，并继续撰写后续章节。"
        )
    },
    {
        "date": "2026年4月15日",
        "content": (
            "检查了论文修改稿的完成情况。学生已按意见完成前三章的修改，第四、五章的实验分析和结论撰写基本完成。"
            "指出实验数据分析部分缺少对异常结果（如JPEG质量因子45处出现的异常恢复）的解释讨论，"
            "建议补充分析说明。参考文献格式需统一调整，正文引用与参考文献列表的编号需逐一核对。"
            "要求学生在4月25日前完成论文定稿并提交查重。"
        )
    },
    {
        "date": "2026年5月1日",
        "content": (
            "学生提交了论文定稿，查重结果符合学校要求（重复率低于20%）。"
            "对最终稿进行了全面审阅，确认论文结构完整、逻辑清晰、实验数据翔实、结论可信。"
            "指导学生准备答辩PPT，建议重点突出研究动机、方法设计和核心实验结果，"
            "预演答辩陈述并准备可能的问题。要求学生继续完善答辩材料，确保答辩顺利进行。"
        )
    },
]

# ── Add extra record rows if needed ──
existing = 6  # template has 6 records
needed = len(records)
if needed > existing:
    # Reference: rows 13-14 are the last existing record (record 6)
    # Each record = 2 rows (time + content)
    # We'll duplicate the last record's rows and insert after the last row
    last_row = tbl.rows[-1]
    ref_elem = last_row._tr

    for i in range(existing, needed):
        # Copy the pattern from the last existing record (rows 13-14)
        # Create 2 new rows for each additional record
        for _ in range(2):
            new_row = copy.deepcopy(tbl.rows[3]._tr)  # copy from record 1 structure
            ref_elem.addnext(new_row)
            ref_elem = new_row

# Re-read table rows after modification
# The table rows object should auto-update
doc.save(OUT)
print(f"Saved intermediate: {OUT}")

# Re-open to fill content
doc2 = Document(OUT)
tbl2 = doc2.tables[0]

# ── Fill each record (starting from row 3, each record = 2 rows) ──
def fill_record(tbl, record_idx, date_text, content_text):
    """Fill record at index record_idx (0-based) with date and content."""
    row_time = 3 + record_idx * 2
    row_content = 4 + record_idx * 2

    if row_content >= len(tbl.rows):
        print(f"Warning: row {row_content} out of range (max {len(tbl.rows)-1})")
        return

    # Fill date in the merged cell (col 1 has gridSpan=5)
    time_cell = tbl.rows[row_time].cells[1]
    set_cell_text(time_cell, f'指导时间：{date_text}')

    # Fill content
    content_cell = tbl.rows[row_content].cells[1]
    full_text = f'指导内容及指导意见：\n{content_text}\n\n指导教师签名：'
    set_cell_text(content_cell, full_text)

for idx, rec in enumerate(records):
    fill_record(tbl2, idx, rec["date"], rec["content"])

doc2.save(OUT)
print(f"\nSaved: {OUT}")
print(f"Filled {len(records)} guidance records")
