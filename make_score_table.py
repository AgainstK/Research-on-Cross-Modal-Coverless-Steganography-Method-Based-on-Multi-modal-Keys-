from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
run = title.add_run('暨南大学本科毕业论文（设计）指导教师评分表')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Student info ──
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.paragraph_format.space_after = Pt(4)
run = info.add_run('学生姓名：张延开    学号：2022101194    专业：网络空间安全')
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.LEFT
info2.paragraph_format.space_after = Pt(6)
run = info2.add_run('论文题目：基于多模态密钥的跨模态无载体隐写方法的研究    指导教师：李佩雅')
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Scoring table ──
# Columns: 评价内容 | 评价要素 | 分值 | 得分 | 评分依据
headers = ['评价内容', '评价要素', '分值', '得分', '评分依据']
rows_data = [
    # (content, element, score, got, basis) — content=None means merged
    # == 选题意义 (15) ==
    ['选题意义\n（15分）',
     '与本专业培养目标和毕业要求的契合程度',
     '5', '5',
     '选题属于网络空间安全专业方向，聚焦扩散模型隐写这一前沿领域，与专业培养目标高度契合。'],
    [None,
     '理论意义或实践应用价值',
     '5', '5',
     '在理论层面深化了对扩散模型潜空间信息隐藏机制的认识，在应用层面为隐蔽通信提供了可行方案，具有一定的学术参考价值和实践意义。'],
    [None,
     '创新意识和难易程度',
     '5', '4',
     '选题体现独立思考和学科研究特点，采用多模态密钥与潜空间调制相结合的思路有一定新意。但跨模态隐写已有相关先例，创新程度尚可进一步提高。难度适中，工作量饱满。'],

    # == 逻辑构建 (25) ==
    ['逻辑构建\n（25分）',
     '结构体例',
     '9', '9',
     '论文核心模块完备，包含绪论、相关技术、方案设计、实验分析、结论等章节，结构完整合理，层次分明，详略得当。'],
    [None,
     '内容组织',
     '8', '8',
     '研究路径清晰，从问题分析到方案设计再到实验验证的逻辑链条完整。实验设计系统全面，涵盖参数敏感性、鲁棒性、对比实验等多个维度，论证充分，结论可信。'],
    [None,
     '文字表达',
     '8', '7',
     '论点表述明确，概念准确，理论运用恰当，论述语言严谨。部分章节的语言表达可进一步精炼，个别表述可更加规范。'],

    # == 专业能力 (40) ==
    ['专业能力\n（40分）',
     '文献检索及梳理能力',
     '10', '10',
     '掌握了文献检索方法，查阅了大量扩散模型与隐写技术相关的中外文献，文献综述部分条理清晰，能按照技术发展脉络进行梳理阐述。'],
    [None,
     '对本专业及相关领域研究现状的了解与评析',
     '10', '9',
     '基本了解本领域学术进展及最新研究动态，对现有方法的优缺点有一定评析能力，能从文献分析中发现问题并提出解决方案。但对部分前沿工作的覆盖可更全面。'],
    [None,
     '对基础理论和专门知识的掌握与运用',
     '10', '10',
     '专业知识扎实，对扩散模型、DDIM采样、隐写技术等核心概念理解准确，理论基础较为扎实，体现出良好的思辨能力和初步的创新能力。'],
    [None,
     '分析和解决问题的能力',
     '10', '10',
     '能够综合运用Python、PyTorch、Diffusers等工具独立完成系统实现，实验中遇到的技术问题（如潜空间分布破坏、两阶段时序控制失当等）均能分析原因并找到解决方案，具备较强的实际问题解决能力。'],

    # == 学术规范 (20) ==
    ['学术规范\n（20分）',
     '论文格式、图表注释、语言规范',
     '10', '10',
     '论文格式符合要求，图表标注规范，中外文用词基本准确，语言通顺。实验数据表格和流程图清晰完整。'],
    [None,
     '写作过程规范性、论文字数',
     '10', '9',
     '论文写作过程合乎规范，相关过程材料完整。论文字数符合要求。参考文献格式存在个别不一致之处，已指出并修改。'],
]

def set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Remove paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)

def set_cell_left(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)

n_rows = len(rows_data) + 1  # +1 for header
tbl = doc.add_table(rows=n_rows, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

# Set column widths
widths = [Cm(2.8), Cm(6.5), Cm(1.5), Cm(1.5), Cm(10.5)]
for row in tbl.rows:
    for ci, w in enumerate(widths):
        row.cells[ci].width = w

# Header
header_texts = ['评价内容', '评价要素', '分值', '得分', '评分依据']
for ci, ht in enumerate(header_texts):
    set_cell(tbl.rows[0].cells[ci], ht, bold=True, size=10)

# Data rows
for ri, rd in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    for ci in range(5):
        val = rd[ci]
        if val is None:
            # merged cell - leave empty
            continue
        if ci == 4:
            set_cell_left(row.cells[ci], val, size=9)
        elif ci == 0:
            set_cell(row.cells[ci], val, bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif ci in [2, 3]:
            set_cell(row.cells[ci], val, bold=True, size=11)
        else:
            set_cell(row.cells[ci], val, size=9.5)

# Merge cells for evaluation content column (col 0)
# 选题意义: rows 1-3 → merge
start_merge = [(1, 3), (4, 6), (7, 10), (11, 12)]
for s, e in start_merge:
    cell_start = tbl.rows[s].cells[0]
    for m in range(s + 1, e + 1):
        cell_start.merge(tbl.rows[m].cells[0])

# ── Total score row ──
total_row = tbl.add_row()
# Merge all cols
total_row.cells[0].merge(total_row.cells[-1])
p = total_row.cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('合计总分：95分（满分100分）')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

# ── Scoring summary ──
doc.add_paragraph()
summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
lines = [
    '评分综述：',
    '该生选题符合网络空间安全专业培养目标，具有一定的理论意义和实践应用价值。论文结构完整、逻辑清晰，从问题分析到方案设计再到实验验证的研究路径合理。',
    '该生掌握了扎实的专业基础知识，能够独立完成系统的算法实现和实验验证，在实验过程中遇到技术问题时能够主动分析原因并找到解决方案，体现出较强的独立科研能力。',
    '论文实验设计较为系统全面，涵盖了参数敏感性分析、鲁棒性测试、对比实验等多个维度，数据翔实、结论可信。论文写作规范，格式符合要求，语言表达通顺。',
    '',
    '不足之处：文献综述对部分前沿工作的覆盖可进一步拓宽；论文部分章节的语言表达可更加精炼；参考文献格式存在个别不一致之处。',
    '',
    '综合评定：该论文达到本科毕业论文水平，建议成绩为优秀（95分）。',
]
for line in lines:
    run = summary.add_run(line + '\n')
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Advisor signature area ──
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig.paragraph_format.space_before = Pt(30)
run = sig.add_run('指导教师签名：李佩雅')
run.font.size = Pt(12)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sig2 = doc.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sig2.add_run('2026年5月6日')
run.font.size = Pt(12)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Save ──
out = '/home/ubuntu/pythonProject1/RESULT/fine_experiments/指导教师评分表.docx'
doc.save(out)
print(f"Saved: {out}")
